#!/usr/bin/env python3
"""Create a sample invoice DOCX for extraction task."""
import sys
from docx import Document
from docx.shared import Pt

doc = Document()
doc.add_heading("Invoice #2026-0042", level=1)
doc.add_paragraph("Date: 2026-02-15")
doc.add_paragraph("From: Acme Corp")
doc.add_paragraph("To: Widget Industries")

doc.add_heading("Line Items", level=2)
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Item"
hdr[1].text = "Quantity"
hdr[2].text = "Unit Price"
hdr[3].text = "Total"

items = [
    ("Web Development", "40", "$150.00", "$6,000.00"),
    ("UI Design", "20", "$120.00", "$2,400.00"),
    ("QA Testing", "15", "$90.00", "$1,350.00"),
    ("Project Management", "10", "$130.00", "$1,300.00"),
    ("DevOps Setup", "8", "$140.00", "$1,120.00"),
]
for item, qty, price, total in items:
    row = table.add_row().cells
    row[0].text = item
    row[1].text = qty
    row[2].text = price
    row[3].text = total

doc.add_paragraph("")
doc.add_paragraph("Subtotal: $12,170.00")
doc.add_paragraph("Tax (8%): $973.60")
doc.add_paragraph("Total Due: $13,143.60")
doc.add_paragraph("")
doc.add_paragraph("Payment Terms: Net 30")
doc.add_paragraph("Due Date: 2026-03-17")

doc.save(sys.argv[1] if len(sys.argv) > 1 else "invoice.docx")
